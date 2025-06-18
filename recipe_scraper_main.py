#get recipe URL from user
recipeurl = input("Enter recipe url: ")
#validate recipe URL & error handling
validrecipieurl = recipeurl.startswith("https://old.reddit.com/r/recipes/comments")
if validrecipieurl == False:
    print('Please enter a URL from https://old.reddit.com/r/recipes')
    while validrecipieurl == False:
        recipeurl = input("Enter recipe url: ")
        validrecipieurl = recipeurl.startswith("https://old.reddit.com/r/recipes/comments")
if validrecipieurl == True:
    print('Processing...')

#GET request to fetch HTML content from validrecipeurl
import requests
from bs4 import BeautifulSoup
recipehtml = requests.get(recipeurl).text

#parse html
recipesoup = BeautifulSoup(recipehtml, 'html.parser')

#Find and Show recipe + comments by div element
div_element = recipesoup.find('div', attrs={'class': 'commentarea'})
div_text = div_element.get_text(separator='\n')
recipetitle = recipesoup.title.text

#Create recipe file
from docx import Document
recipedoc = Document()
#write recipe to file
recipedoc.add_heading(recipetitle, level=1)
recipedoc.add_paragraph(div_text)
#define filename
user_input = input("Enter filename (without extension): ")
filename = f"{user_input.strip()}.docx"
#save the file
recipedoc.save(filename)
print("File saved to", filename)
